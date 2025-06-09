"""
Document Processing Module for Tender Documents
Handles PDF, Word, and other document formats for AI analysis
"""
import os
import time
import requests
import logging
from typing import Dict, List, Optional, Tuple
from urllib.parse import urljoin, urlparse
import hashlib
from datetime import datetime
import re
from models import TenderDocument, SessionLocal
from ai_analyzer import DocumentAnalyzer

# Try to import document processing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("PyPDF2 not available - PDF processing disabled")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available - Word processing disabled")

from io import BytesIO

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process and analyze tender documents"""
    
    def __init__(self):
        """Initialize document processor"""
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        self.analyzer = DocumentAnalyzer()
        
        # Create downloads directory
        self.download_dir = "c:\\Users\\user\\Projects\\Scraping\\downloads"
        os.makedirs(self.download_dir, exist_ok=True)
    
    def download_document(self, url: str, tender_id: int = None) -> Optional[str]:
        """Download document from URL"""
        try:
            # Get filename from URL
            parsed_url = urlparse(url)
            filename = os.path.basename(parsed_url.path)
            
            # If no filename, generate one
            if not filename or '.' not in filename:
                url_hash = hashlib.md5(url.encode()).hexdigest()[:8]
                filename = f"document_{url_hash}.pdf"
            
            # Create full path
            filepath = os.path.join(self.download_dir, filename)
            
            # Download with timeout and size limit
            response = self.session.get(url, timeout=30, stream=True)
            response.raise_for_status()
            
            # Check content type
            content_type = response.headers.get('content-type', '').lower()
            if 'pdf' not in content_type and 'word' not in content_type and 'document' not in content_type:
                logger.warning(f"Unexpected content type for {url}: {content_type}")
            
            # Check file size (limit to 50MB)
            content_length = response.headers.get('content-length')
            if content_length and int(content_length) > 50 * 1024 * 1024:
                logger.warning(f"File too large: {url} ({content_length} bytes)")
                return None
            
            # Download in chunks
            with open(filepath, 'wb') as f:
                downloaded = 0
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
                        downloaded += len(chunk)
                        # Size limit check during download
                        if downloaded > 50 * 1024 * 1024:
                            logger.warning(f"File too large during download: {url}")
                            os.remove(filepath)
                            return None
            
            logger.info(f"Downloaded document: {filename} ({downloaded} bytes)")
            return filepath
            
        except Exception as e:
            logger.error(f"Error downloading document {url}: {str(e)}")
            return None
    
    def extract_text_from_pdf(self, filepath: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Limit to first 20 pages
                max_pages = min(len(pdf_reader.pages), 20)
                
                for page_num in range(max_pages):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\\n"
            
            # Clean up text
            text = re.sub(r'\\s+', ' ', text)  # Multiple whitespace to single space
            text = re.sub(r'\\n+', '\\n', text)  # Multiple newlines to single
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {filepath}: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, filepath: str) -> str:
        """Extract text from Word document"""
        try:
            doc = docx.Document(filepath)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\\n"
            
            # Clean up text
            text = re.sub(r'\\s+', ' ', text)
            text = re.sub(r'\\n+', '\\n', text)
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {filepath}: {str(e)}")
            return ""
    
    def extract_text_from_document(self, filepath: str) -> str:
        """Extract text from document based on file type"""
        try:
            filename = os.path.basename(filepath).lower()
            
            if filename.endswith('.pdf'):
                return self.extract_text_from_pdf(filepath)
            elif filename.endswith(('.docx', '.doc')):
                return self.extract_text_from_docx(filepath)
            else:
                # Try to read as plain text
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
        except Exception as e:
            logger.error(f"Error extracting text from {filepath}: {str(e)}")
            return ""
    
    def analyze_document_content(self, text: str, document_url: str) -> Dict:
        """Analyze document content using AI"""
        try:
            if not text or len(text.strip()) < 100:
                logger.warning(f"Document text too short for analysis: {document_url}")
                return {}
            
            # Use AI analyzer
            analysis = self.analyzer.analyze_document_content(text)
            
            # Extract key information
            analysis_result = {
                'key_requirements': self.extract_requirements(text),
                'deadline_mentions': self.extract_deadlines(text),
                'budget_mentions': self.extract_budget_info(text),
                'location_mentions': self.extract_locations(text),
                'sector_indicators': self.extract_sector_info(text),
                'ai_analysis': analysis
            }
            
            return analysis_result
            
        except Exception as e:
            logger.error(f"Error analyzing document content: {str(e)}")
            return {}
    
    def extract_requirements(self, text: str) -> List[str]:
        """Extract key requirements from text"""
        requirements = []
        text_lower = text.lower()
        
        # Look for requirement indicators
        requirement_patterns = [
            r'requirement[s]?:?\\s*([^.\\n]{20,100})',
            r'must\\s+([^.\\n]{20,100})',
            r'shall\\s+([^.\\n]{20,100})',
            r'mandatory:?\\s*([^.\\n]{20,100})',
            r'essential:?\\s*([^.\\n]{20,100})'
        ]
        
        for pattern in requirement_patterns:
            matches = re.findall(pattern, text_lower, re.IGNORECASE)
            requirements.extend([match.strip() for match in matches[:5]])  # Limit to 5 per pattern
        
        return requirements[:10]  # Limit total requirements
    
    def extract_deadlines(self, text: str) -> List[str]:
        """Extract deadline mentions from text"""
        deadlines = []
        
        # Date patterns
        date_patterns = [
            r'deadline:?\\s*([^.\\n]{10,50})',
            r'due\\s+(?:date|by):?\\s*([^.\\n]{10,50})',
            r'submission\\s+(?:date|deadline):?\\s*([^.\\n]{10,50})',
            r'close[s]?\\s+(?:on|at):?\\s*([^.\\n]{10,50})'
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            deadlines.extend([match.strip() for match in matches[:3]])
        
        return deadlines[:5]
    
    def extract_budget_info(self, text: str) -> List[str]:
        """Extract budget information from text"""
        budget_info = []
        
        # Budget patterns
        budget_patterns = [
            r'budget:?\\s*([^.\\n]{10,50})',
            r'value:?\\s*([£$€][^.\\n]{5,30})',
            r'worth:?\\s*([£$€][^.\\n]{5,30})',
            r'contract\\s+value:?\\s*([^.\\n]{10,50})'
        ]
        
        for pattern in budget_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            budget_info.extend([match.strip() for match in matches[:3]])
        
        return budget_info[:5]
    
    def extract_locations(self, text: str) -> List[str]:
        """Extract location mentions from text"""
        locations = []
        
        # Common locations for Square Circle
        location_keywords = [
            'pacific', 'fiji', 'vanuatu', 'tonga', 'samoa', 'solomon islands',
            'papua new guinea', 'png', 'australia', 'new zealand', 'kiribati',
            'marshall islands', 'micronesia', 'palau', 'nauru', 'tuvalu'
        ]
        
        text_lower = text.lower()
        for keyword in location_keywords:
            if keyword in text_lower:
                # Find context around the keyword
                pattern = f'.{{0,30}}{re.escape(keyword)}.{{0,30}}'
                matches = re.findall(pattern, text_lower)
                locations.extend([match.strip() for match in matches[:2]])
        
        return locations[:10]
    
    def extract_sector_info(self, text: str) -> List[str]:
        """Extract sector information from text"""
        sectors = []
        
        # Sector keywords relevant to Square Circle
        sector_keywords = [
            'climate', 'environment', 'governance', 'infrastructure',
            'development', 'capacity building', 'policy', 'planning',
            'sustainability', 'resilience', 'adaptation', 'mitigation'
        ]
        
        text_lower = text.lower()
        for keyword in sector_keywords:
            if keyword in text_lower:
                # Find context around the keyword
                pattern = f'.{{0,40}}{re.escape(keyword)}.{{0,40}}'
                matches = re.findall(pattern, text_lower)
                sectors.extend([match.strip() for match in matches[:2]])
        
        return sectors[:10]
    
    def process_document_url(self, url: str, tender_id: int = None) -> Optional[Dict]:
        """Process a document URL completely"""
        try:
            logger.info(f"Processing document: {url}")
            
            # Download document
            filepath = self.download_document(url, tender_id)
            if not filepath:
                return None
            
            # Extract text
            text = self.extract_text_from_document(filepath)
            if not text:
                logger.warning(f"No text extracted from {url}")
                # Clean up file
                try:
                    os.remove(filepath)
                except:
                    pass
                return None
            
            # Analyze content
            analysis = self.analyze_document_content(text, url)
            
            # Prepare result
            result = {
                'url': url,
                'filepath': filepath,
                'text_length': len(text),
                'text_preview': text[:500] + "..." if len(text) > 500 else text,
                'analysis': analysis,
                'processed_at': datetime.utcnow()
            }
            
            # Save to database if tender_id provided
            if tender_id:
                self.save_document_to_db(result, tender_id)
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {url}: {str(e)}")
            return None
    
    def save_document_to_db(self, doc_data: Dict, tender_id: int):
        """Save document data to database"""
        db = SessionLocal()
        try:
            # Check if document already exists
            existing = db.query(TenderDocument).filter_by(
                url=doc_data['url'],
                tender_id=tender_id
            ).first()
            
            if not existing:
                doc = TenderDocument(
                    tender_id=tender_id,
                    url=doc_data['url'],
                    file_path=doc_data['filepath'],
                    text_content=doc_data['text_preview'],
                    analysis_data=str(doc_data['analysis']),
                    processed_at=doc_data['processed_at']
                )
                db.add(doc)
                db.commit()
                logger.info(f"Saved document to database: {doc_data['url']}")
            
        except Exception as e:
            logger.error(f"Error saving document to database: {str(e)}")
            db.rollback()
        finally:
            db.close()
    
    def process_tetra_tech_documents(self) -> List[Dict]:
        """Process documents from Tetra Tech that were found earlier"""
        # URLs found from earlier scraping
        tetra_tech_docs = [
            "https://www.tetratechintdev.com/pdfs/AIFFP_AM-AIFFP-2024-001_RTF_FINAL.pdf",
            "https://www.tetratechintdev.com/pdfs/SC_Public_Consultation_Strategy_and_Plan.pdf",
            "https://www.tetratechintdev.com/pdfs/SC_Youth_Development_Strategy.pdf"
        ]
        
        results = []
        for doc_url in tetra_tech_docs:
            try:
                result = self.process_document_url(doc_url)
                if result:
                    results.append(result)
                    logger.info(f"Processed Tetra Tech document: {doc_url}")
                
                # Wait between downloads
                time.sleep(2)
                
            except Exception as e:
                logger.error(f"Error processing Tetra Tech document {doc_url}: {str(e)}")
                continue
        
        return results

if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Test with Tetra Tech documents
    print("Processing Tetra Tech documents...")
    results = processor.process_tetra_tech_documents()
    print(f"Processed {len(results)} documents")
    
    for result in results:
        print(f"\\nDocument: {result['url']}")
        print(f"Text length: {result['text_length']} characters")
        print(f"Preview: {result['text_preview'][:200]}...")
        if result['analysis']:
            print(f"Analysis keys: {list(result['analysis'].keys())}")
